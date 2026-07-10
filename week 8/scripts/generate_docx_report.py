import os
import sys

# Define path constants
MEDIA_DIR = r"C:\Users\Manoj Pareek\.gemini\antigravity-ide\brain\8b9a75ba-9c00-45d1-babb-62ca83790ec6"
OUTPUT_PATH = r"c:\Users\Manoj Pareek\OneDrive\Desktop\bcccc\output\Ecommerce_Project_Report.docx"

def create_report():
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Error: python-docx is not installed. Please run 'pip install python-docx' first.")
        sys.exit(1)
        
    doc = Document()
    
    # Configure document margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("E-Commerce Order Analytics System")
    title_run.font.size = Pt(24)
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Project Implementation & Execution Report")
    subtitle_run.font.size = Pt(14)
    subtitle_run.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("\n")
    
    # Step definitions
    steps = [
        {
            "num": "1",
            "title": "Mock Data Generation",
            "desc": "We run 'generate_data.py' to generate raw customer, product, order, and order item CSV files with simulated data anomalies.",
            "img": "media__1783668459837.png"
        },
        {
            "num": "2",
            "title": "Programmatic Data Cleaning",
            "desc": "We clean the raw CSV files using Pandas, standardizing formats, and validating relational constraints.",
            "img": "media__1783668442454.png"
        },
        {
            "num": "3",
            "title": "Edge Case Testing and Ingestion",
            "desc": "We verify validation rules using unit tests, deploy the database schema in SQLite, and import the cleaned datasets.",
            "img": "media__1783668459900.png"
        },
        {
            "num": "4",
            "title": "Execute SQL Aggregations & Joins",
            "desc": "We run basic and intermediate SQL queries to analyze category revenues, top customers, order volumes, and return rates.",
            "img": "media__1783668459849.png"
        },
        {
            "num": "5",
            "title": "Execute SQL Window Functions & CTEs",
            "desc": "We run advanced SQL window functions and multi-level CTEs to analyze product rankings, spend gaps, and cohort metrics.",
            "img": "media__1783668417380.png"
        }
    ]
    
    for step in steps:
        h = doc.add_paragraph()
        h_run = h.add_run(f"Step {step['num']}: {step['title']}")
        h_run.font.size = Pt(16)
        h_run.bold = True
        
        d = doc.add_paragraph()
        d_run = d.add_run(step['desc'])
        d_run.font.size = Pt(11)
        
        img_path = os.path.join(MEDIA_DIR, step['img'])
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(6.2))
        else:
            doc.add_paragraph(f"[Screenshot not found: {step['img']}]")
            
        doc.add_paragraph("\n")
        
    doc.save(OUTPUT_PATH)
    print(f"Report successfully saved to {OUTPUT_PATH}")

if __name__ == "__main__":
    create_report()
